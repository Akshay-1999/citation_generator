import datetime
from pathlib import Path
from typing import AsyncIterator, List, Optional, Tuple
import asyncio
import aiofiles
import io
import re,os
import os
from bs4 import BeautifulSoup
import fitz
import pymupdf4llm
from langchain.schema import Document
import time
from utils.logging_utils import set_system_logger
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
logger = set_system_logger("system_logger")


async def extract_with_pymupdf(file_path: str, page_range: Optional[List[int]] = None) -> str:
    start_time = time.perf_counter()
    start_timestamp = datetime.datetime.now()
    logger.info(f"=== USING PYMUPDF TOOL FOR FILE: {Path(file_path).name} ===")
    metadata={
        "chunk_name": file_path,
        "processing_tool": 'pymupdf4llm',
        "start_timestamp": start_timestamp,
        "processing_time": 0,
        "task_id": "N/A",
        "fallback_reason": "N/A"
    }
    loop = asyncio.get_running_loop()
    try:
        if page_range:
            result = await loop.run_in_executor(
                None, 
                lambda: pymupdf4llm.to_markdown(file_path, pages=page_range)
            ) or ""
        else:
            result = await loop.run_in_executor(
                None, 
                lambda: pymupdf4llm.to_markdown(file_path)
            ) or ""
            
        # Detect silent failures where pymupdf4llm only extracts tables and misses the main text
        doc = fitz.open(file_path)
        try:
            fitz_len = sum(len(page.get_text()) for page in (doc[i] for i in (page_range if page_range else range(len(doc)))))
            if len(result.strip()) < 1000 and fitz_len > len(result.strip()) + 500:
                raise Exception(f"pymupdf4llm returned only {len(result.strip())} chars, but fitz found {fitz_len} chars. Forcing fallback.")
            
            # Prevent resume headers (Name, Contact) from being stripped
            if len(doc) > 0:
                first_page_text = doc[0].get_text().strip()
                if first_page_text:
                    import re
                    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
                    raw_emails = set(re.findall(email_pattern, first_page_text.lower()))
                    md_emails = set(re.findall(email_pattern, result.lower()))
                    
                    # Basic phone pattern
                    phone_pattern = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
                    raw_phones = set(re.findall(phone_pattern, first_page_text))
                    md_phones = set(re.findall(phone_pattern, result))
                    
                    header_text = first_page_text[:300]
                    # If the first 50 characters of the raw text are missing from the markdown, or if contact info is missing
                    heuristic_miss = bool(header_text[:50].strip() and header_text[:50].strip() not in result)
                    missing_email = bool(len(raw_emails - md_emails) > 0)
                    missing_phone = bool(len(raw_phones - md_phones) > 0)
                    
                    if heuristic_miss or missing_email or missing_phone:
                        logger.info(f"--- Prepended stripped header for {Path(file_path).name} (heuristic_miss={heuristic_miss}, missing_email={missing_email}, missing_phone={missing_phone}) ---")
                        result = "--- RAW TEXT FALLBACK FOR MISSING HEADER INFO ---\n" + first_page_text + "\n--- END RAW TEXT FALLBACK ---\n\n" + result
        finally:
            doc.close()

        processing_time = time.perf_counter() - start_time
        metadata["processing_time"] = processing_time
        logger.info(f"=== PYMUPDF EXTRACTION COMPLETED FOR FILE: {Path(file_path).name} - SUCCESS IN {processing_time:.2f} seconds ===")
        return result, metadata
    except Exception as e:
        logger.error(f"=== pymupdf4llm extraction failed: {e} ===")
        # Fallback: extract plain text using fitz
        logger.info(f"--- Falling back to fitz plain-text extraction for: {Path(file_path).name} ---")
        try:
            def _fitz_fallback():
                doc = None
                try:
                    doc = fitz.open(file_path)
                    pages = page_range if page_range else range(len(doc))
                    text_parts = []
                    for page_num in pages:
                        if page_num < len(doc):
                            text_parts.append(doc[page_num].get_text())
                    return "\n\n".join(text_parts)
                finally:
                    if doc:
                        doc.close()

            result = await loop.run_in_executor(None, _fitz_fallback)
            metadata["processing_tool"] = "fitz_fallback"
            metadata["fallback_reason"] = str(e)
            processing_time = time.perf_counter() - start_time
            metadata["processing_time"] = processing_time
            logger.info(f"=== FITZ FALLBACK COMPLETED FOR FILE: {Path(file_path).name} - {'SUCCESS' if result.strip() else 'NO TEXT'} IN {processing_time:.2f} seconds ===")
            return result, metadata
        except Exception as fallback_err:
            logger.error(f"=== Fitz fallback also failed for {Path(file_path).name}: {fallback_err} ===")
            processing_time = time.perf_counter() - start_time
            metadata["processing_time"] = processing_time
            return "", metadata

def _is_html_disguised_as_doc(file_path: str) -> bool:
    """
    Naukri.com (and some other portals) download resumes as HTML files
    but save them with a .doc extension. Detect this by sniffing the
    first 512 bytes for an HTML signature.
    """
    try:
        with open(file_path, "rb") as f:
            header = f.read(512).lstrip()  # skip any leading whitespace/BOM
        header_lower = header[:100].lower()
        return (
            header.startswith(b"<!doctype html") or
            header.startswith(b"<!DOCTYPE html") or
            b"<html" in header_lower or
            b"<html>" in header_lower
        )
    except Exception as e:
        logger.warning(f"--- _is_html_disguised_as_doc could not read {file_path}: {e} ---")
        return False


def _extract_text_from_html_doc(file_path: str) -> str:
    """
    Extract visible text from a .doc file that is actually HTML (e.g. Naukri downloads).
    Uses BeautifulSoup for clean, dependency-free extraction.
    """
    # Detect encoding — try UTF-8 first, then latin-1 as safe fallback
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=enc, errors="replace") as f:
                raw_html = f.read()
            break
        except Exception:
            continue
    else:
        return ""

    soup = BeautifulSoup(raw_html, "html.parser")

    # Remove script and style elements that pollute text output
    for tag in soup(["script", "style", "head", "meta", "link"]):
        tag.decompose()

    # Extract text, preserving paragraph breaks
    lines = []
    for element in soup.find_all(["p", "div", "li", "td", "th", "h1", "h2", "h3", "h4", "h5", "h6", "span", "br"]):
        text = element.get_text(separator=" ", strip=True)
        if text:
            lines.append(text)

    # Deduplicate consecutive identical lines (tables sometimes repeat)
    deduped = []
    prev = None
    for line in lines:
        if line != prev:
            deduped.append(line)
        prev = line

    return "\n\n".join(deduped)


async def extract_with_unstructured(file_path: str, page_range: Optional[List[int]] = None) -> str:
    start_time = time.perf_counter()
    start_timestamp = datetime.datetime.now()
    logger.info(f"=== USING UNSTRUCTURED TOOL FOR FILE: {Path(file_path).name} ===")
    metadata={
        "chunk_name": file_path,
        "processing_tool": 'unstructured_docx',
        "start_timestamp": start_timestamp,
        "processing_time": 0,
        "task_id": "N/A",
        "fallback_reason": "N/A",
        "number_of_pages": None
    }
    loop = asyncio.get_running_loop()

    # ── STEP 1: detect HTML-disguised .doc files (common with Naukri.com downloads) ──
    try:
        if _is_html_disguised_as_doc(file_path):
            logger.info(f"--- Detected HTML-disguised .doc file: {Path(file_path).name}. Using BeautifulSoup extractor ---")
            result = await loop.run_in_executor(None, _extract_text_from_html_doc, file_path)
            processing_time = time.perf_counter() - start_time
            metadata["processing_tool"] = "beautifulsoup_html_doc"
            metadata["processing_time"] = processing_time
            if result and result.strip():
                logger.info(f"=== HTML-DOC EXTRACTION COMPLETED FOR FILE: {Path(file_path).name} - {len(result)} chars IN {processing_time:.2f} seconds ===")
                return result, metadata
            else:
                logger.warning(f"=== BeautifulSoup extracted empty text from HTML-disguised doc: {Path(file_path).name} ===")
    except Exception as html_err:
        logger.error(f"=== HTML-doc extraction failed for {Path(file_path).name}: {html_err} ===")

    # ── STEP 2: standard unstructured loader for real .docx / OLE .doc files ──
    try:
        def _load_docx():
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            docs = loader.load()
            # Combine all elements into a single text string
            text = "\n\n".join([doc.page_content for doc in docs if doc.page_content.strip()])
            return text
        
        result = await loop.run_in_executor(None, _load_docx)
        
        processing_time = time.perf_counter() - start_time
        metadata["processing_time"] = processing_time
        
        logger.info(f"=== UNSTRUCTURED EXTRACTION COMPLETED FOR FILE: {Path(file_path).name} - SUCCESS IN {processing_time:.2f} seconds ===")
        
        return result, metadata

    except Exception as e:
        logger.warning(f"=== unstructured loader failed for {Path(file_path).name}: {e}. Trying python-docx fallback ===")
        
    # ── STEP 3: python-docx fallback for .docx files ──
    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            def _load_with_python_docx():
                import docx
                doc = docx.Document(file_path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                # Also extract tables
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            table_texts.append(row_text)
                return "\n\n".join(paragraphs + table_texts)

            result = await loop.run_in_executor(None, _load_with_python_docx)
            processing_time = time.perf_counter() - start_time
            metadata["processing_tool"] = "python_docx_fallback"
            metadata["processing_time"] = processing_time
            if result and result.strip():
                logger.info(f"=== python-docx FALLBACK COMPLETED FOR FILE: {Path(file_path).name} - SUCCESS IN {processing_time:.2f} seconds ===")
                return result, metadata
    except Exception as docx_err:
        logger.error(f"=== python-docx fallback also failed for {Path(file_path).name}: {docx_err} ===")

    # ── STEP 4: olefile raw stream extractor for true OLE binary .doc files (Word 97-2003) ──
    # This handles real .doc compound document files without needing LibreOffice.
    try:
        def _extract_from_ole_stream(fp: str) -> str:
            import olefile
            import re as _re
            ole = olefile.OleFileIO(fp)
            try:
                if not ole.exists('WordDocument'):
                    return ""
                raw = ole.openstream('WordDocument').read()
                # Word 97-2003 stores text as UTF-16LE.
                # The first 768 bytes (384 UTF-16LE chars) are the FIB binary header — skip them.
                decoded = raw.decode('utf-16-le', errors='replace')[384:]
                # Keep printable ASCII, common Latin/Unicode letters and whitespace
                cleaned = _re.sub(r'[^\x20-\x7E\x09\x0A\x0D\u00A0-\u2FFF]', ' ', decoded)
                cleaned = _re.sub(r'  +', ' ', cleaned)
                # Split on control chars / null bytes and keep lines with real content
                raw_lines = _re.split(r'[\x00\x01-\x08\x0B\x0C\x0E-\x1F\r\n]+', cleaned)
                lines = []
                for l in raw_lines:
                    l = l.strip()
                    if len(l) < 4:
                        continue
                    # Filter out binary-header remnants: lines where <40% chars are alpha/space
                    alpha_ratio = sum(c.isalpha() or c.isspace() for c in l) / len(l)
                    if alpha_ratio < 0.40:
                        continue
                    # Strip any leading garbage tokens: sequences of ? and space-separated single chars
                    # e.g. "n 0 h ? ? ? NAGARAJU" → "NAGARAJU"
                    l = _re.sub(r'^([\?\s\x00-\x1F\d]{1,3}\s+){2,}', '', l).strip()
                    if len(l) < 4:
                        continue
                    lines.append(l)
                # Deduplicate consecutive identical lines
                deduped = []
                prev = None
                for line in lines:
                    if line != prev:
                        deduped.append(line)
                    prev = line
                return "\n".join(deduped)
            finally:
                ole.close()

        logger.info(f"--- Attempting OLE binary stream extraction for: {Path(file_path).name} ---")
        result = await loop.run_in_executor(None, _extract_from_ole_stream, file_path)
        processing_time = time.perf_counter() - start_time
        metadata["processing_tool"] = "olefile_doc_extractor"
        metadata["processing_time"] = processing_time
        if result and result.strip():
            logger.info(f"=== OLE EXTRACTOR COMPLETED FOR FILE: {Path(file_path).name} - {len(result)} chars IN {processing_time:.2f} seconds ===")
            return result, metadata
        else:
            logger.warning(f"=== OLE extractor returned empty text for: {Path(file_path).name} ===")
    except ImportError:
        logger.warning("=== olefile not installed; skipping OLE binary extraction step ===")
    except Exception as ole_err:
        logger.error(f"=== OLE binary extractor failed for {Path(file_path).name}: {ole_err} ===")

    processing_time = time.perf_counter() - start_time
    metadata["processing_time"] = processing_time
    return "", metadata

